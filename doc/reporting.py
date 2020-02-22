#-*- coding:utf-8 -*-
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn 
from docx.shared import Pt
import personal as ps

document = Document()

h = document.add_heading()
hh = h.add_run('Report',0)
hh.bold = True
hh.italic = True
font0 = hh.font
font0.size = Pt(30)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER

p0 = document.add_paragraph()
pp0= p0.add_run('신상정보\n')
pp0.bold = True
pp0.italic = True
font0 = pp0.font
font0.size = Pt(17)

table = document.add_table(rows = 2, cols = 4)
table.style = document.styles['Table Grid']
 
hdr_cells = table.rows[0].cells
hdr_cells[0].paragraphs[0].add_run('이름').bold = True
hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
hdr_cells[1].paragraphs[0].add_run(ps.name1).bold = True
hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[2].paragraphs[0].add_run('성별').bold = True
hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells[3].paragraphs[0].add_run(ps.sex).bold = True
hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells1 = table.rows[1].cells
hdr_cells1[0].paragraphs[0].add_run('전화번호').bold = True
hdr_cells1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
hdr_cells1[1].paragraphs[0].add_run(ps.phone).bold = True
hdr_cells1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells1[2].paragraphs[0].add_run('점수').bold = True
hdr_cells1[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr_cells1[3].paragraphs[0].add_run(ps.score).bold = True
hdr_cells1[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

p1 = document.add_paragraph()
pp1= p1.add_run('\n약관 요약(넣을꺼없어서 넣었음)\n')
pp1.bold = True
pp1.italic = True
font1 = pp1.font
font1.size = Pt(17)

f1 = open('result2.txt', 'rt', encoding='utf-8-sig')
data = f1.readlines()
f1.close()
for datum in data:
    if str(datum):
        pp2 = p1.add_run(datum)
        font2 = pp2.font
        font2.size = Pt(12)
    

p3 = document.add_paragraph()
pp3= p3.add_run('\n얼굴 분석 결과\n')
pp3.bold = True
pp3.italic = True
font3 = pp3.font
font3.size = Pt(17)

p4 = document.add_paragraph()
pp4= p4.add_run('\n녹취록\n')
pp4.bold = True
pp4.italic = True
font4 = pp4.font
font4.size = Pt(17)

document.save('report.docx')