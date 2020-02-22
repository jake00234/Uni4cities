# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'recordwindow.ui'
#
# Created by: PyQt5 UI code generator 5.14.1
#
# WARNING! All changes made in this file will be lost!

endpoint = "https://kpmg-text.cognitiveservices.azure.com/"


from PyQt5 import QtCore, QtGui, QtWidgets,QtTest
from stt import SpeechToText
from Face import EmotionFace
from azure.ai.textanalytics import TextAnalyticsClient, TextAnalyticsApiKeyCredential
from capture import Capture
import sys
import key
key = key.text_key
class Ui_recordWindow(object):

    All=''
    numclicked = 0
    Answer1 = ''
    Answer2 = ''
    Answer3 = ''
    
    

    def __init__(self):
        print("0")
        file=open('questions/question1.txt','r',encoding='utf-8-sig')
        self.All=file.read()
        file.close()
        
        
        
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(800, 600)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.textBrowser = QtWidgets.QTextBrowser(self.centralwidget)
        self.textBrowser.setGeometry(QtCore.QRect(40, 70, 721, 192))
        self.textBrowser.setObjectName("textBrowser")
        self.textBrowser.setText(self.All)
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(40, 10, 161, 61))
        self.label.setObjectName("label")
        self.recordButton = QtWidgets.QPushButton(self.centralwidget)
        self.recordButton.setGeometry(QtCore.QRect(40, 280, 111, 41))
        self.recordButton.setObjectName("recordButton")
        self.recordButton.clicked.connect(self.onrecordButtonClicked)

        self.NextButton = QtWidgets.QPushButton(self.centralwidget)
        self.NextButton.setGeometry(QtCore.QRect(650, 280, 111, 41))
        self.NextButton.setObjectName("NextButton")
        self.NextButton.clicked.connect(self.onNextButtonClicked)

        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 800, 22))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "비대면상품ai서비스"))
        self.label.setText(_translate("MainWindow", "질문"))
        self.recordButton.setText(_translate("MainWindow", "녹음"))
        self.NextButton.setText(_translate("MainWindow", "다음"))

    def onrecordButtonClicked(self):
        print('record')
        if self.numclicked == 0:
            msg=QtWidgets.QMessageBox()
            msg.setWindowTitle('recording')
            msg.setText('recording')
            msg.show()
            stt = SpeechToText()
            img_name = "image{}".format(self.numclicked)
            img = Capture(img_name)
            self.Answer1 = stt.text
            msg.close()
        elif self.numclicked == 1:
            msg=QtWidgets.QMessageBox()
            msg.setWindowTitle('recording')
            msg.setText('recording')
            msg.show()
            stt = SpeechToText()
            img_name = "image{}".format(self.numclicked)
            img = Capture(img_name)
            self.Answer2 = stt.text
            msg.close()
        elif self.numclicked == 2:
            msg=QtWidgets.QMessageBox()
            msg.setWindowTitle('recording')
            msg.setText('recording')
            msg.show()
            stt = SpeechToText()
            img_name = "image{}".format(self.numclicked)
            img = Capture(img_name)
            self.Answer3 = stt.text
            msg.close()
        

    def onNextButtonClicked(self):
        if self.numclicked == 0:
            file=open('./questions/question2.txt','r', encoding='utf-8-sig')
            self.All=file.read()
            file.close()
            self.textBrowser.setText(self.All)
            self.numclicked = self.numclicked + 1

        elif self.numclicked == 1:
            file=open('./questions/question3.txt','r', encoding='utf-8-sig')
            self.All=file.read()
            file.close()
            self.textBrowser.setText(self.All)
            self.numclicked = self.numclicked + 1

        elif self.numclicked == 2:
            self.displayResult()
            self.numclicked = self.numclicked + 1

        elif self.numclicked == 3:
            msg=QtWidgets.QMessageBox()
            msg.setWindowTitle('종료메세지')
            msg.setText('수고하셨습니다. 심사결과는 잠시 후\n문자메시지로 발송됩니다.\n감사합니다.')
            msg.show()
            QtTest.QTest.qWait(3000)
            msg.close()
            sys.exit()



    def displayResult(self):
        text_analytics_client = TextAnalyticsClient(endpoint, TextAnalyticsApiKeyCredential(key))
        documents = [
            self.Answer1,
            self.Answer2,
            self.Answer3,
            
        ]

        response = text_analytics_client.analyze_sentiment(documents, language="ko")
        result = [doc for doc in response if not doc.is_error]
        voice = []
        for doc in result:
            print("Overall sentiment: {}".format(doc.sentiment))
            print("Scores: positive={0:.3f}; neutral={1:.3f}; negative={2:.3f} \n".format(
                doc.sentiment_scores.positive,
                doc.sentiment_scores.neutral,
                doc.sentiment_scores.negative,
            ))
            
            voice.append([doc.sentiment_scores.positive, doc.sentiment_scores.neutral, doc.sentiment_scores.negative])
        print("<----sentiment\n")



        response = text_analytics_client.extract_key_phrases(documents, language="ko")
        result = [doc for doc in response if not doc.is_error]
        answer_sheet = []
        for i in range(3):
            f1 = open('answers/answer'+str(i+1)+'.txt', 'rt', encoding='utf-8-sig')
            answer= f1.read()
            answer_sheet.append(answer)
            f1.close()
        que = 0
        tf = 0
        for doc in result:
            print(doc.key_phrases)
            if doc.key_phrases:
                print(answer_sheet[que])
                abcd = ''.join(doc.key_phrases)
                if answer_sheet[que] in abcd:
                    tf += 1
                    print(tf)
            else:
                print(answer_sheet[que])
                abcdef = ''.join(self.Answer1)
                if answer_sheet[que] in abcdef:
                    tf += 1
                    print(tf)
            que += 1
        result123 = round(tf/que *100,1)
        print('score='+str(result123)+'%')
        print("<----keyphrase\n")

        i = 0
        face_list = []
        while i < 3:
            img_name = "image{}.png".format(i)
            emotionresult = EmotionFace(img_name).emotions
            print('문제'+str(i)+emotionresult)
            face_list.append(emotionresult)
            i= i+1

        import docx
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn 
        from docx.shared import Pt
        import personal as ps
        from docx.shared import Cm, Inches

        document = Document()

        h = document.add_heading()
        hh = h.add_run('Report',0)
        hh.bold = True
        hh.italic = True
        font0 = hh.font
        font0.size = Pt(30)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p0 = document.add_paragraph()
        pp0= p0.add_run('신상정보\n')
        pp0.bold = True
        pp0.italic = True
        font0 = pp0.font
        font0.size = Pt(17)

        table = document.add_table(rows = 2, cols = 4)
        table.style = document.styles['Table Grid']
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run('이름').bold = True
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
        hdr_cells[1].paragraphs[0].add_run(ps.name1).bold = True
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[2].paragraphs[0].add_run('성별').bold = True
        hdr_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells[3].paragraphs[0].add_run(ps.sex).bold = True
        hdr_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells1 = table.rows[1].cells
        hdr_cells1[0].paragraphs[0].add_run('전화번호').bold = True
        hdr_cells1[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
        hdr_cells1[1].paragraphs[0].add_run(ps.phone).bold = True
        hdr_cells1[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells1[2].paragraphs[0].add_run('점수').bold = True
        hdr_cells1[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells1[3].paragraphs[0].add_run(str(result123)).bold = True
        hdr_cells1[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        '''p1 = document.add_paragraph()
        pp1= p1.add_run('\n약관 요약(넣을꺼없어서 넣었음)\n')
        pp1.bold = True
        pp1.italic = True
        font1 = pp1.font
        font1.size = Pt(17)'''

        '''f1 = open('result2.txt', 'rt', encoding='utf-8-sig')
        data = f1.readlines()
        f1.close()
        for datum in data:
            if str(datum):
                pp2 = p1.add_run(datum)
                font2 = pp2.font
                font2.size = Pt(12)
            '''

        p3 = document.add_paragraph()
        pp3= p3.add_run('\n얼굴 분석 결과\n')
        pp3.bold = True
        pp3.italic = True
        font3 = pp3.font
        font3.size = Pt(17)
        #해야됨
        for i in range(3):
            image_path = 'image'+str(i)+'.png'
            document.add_picture(image_path, width= Cm(4.91), height= Cm(8))
            table = document.add_table(rows =1, cols = 1)
            table.style = document.styles['Table Grid']
            new_face = face_list[i].split(', ')
            cell = table.cell(0,0)
            cell.text = face_list[i]
            
    
              
                


        p6 = document.add_paragraph()
        pp6= p6.add_run('\n음성 분석 결과\n')
        pp6.bold = True
        pp6.italic = True
        font6 = pp6.font
        font6.size = Pt(17)

        table3 = document.add_table(rows = 4, cols = 3)
        table3.style = document.styles['Table Grid']
        
        hdr_cells3= table3.rows[0].cells
        hdr_cells3[0].paragraphs[0].add_run('긍정').bold = True
        hdr_cells3[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 
        hdr_cells3[1].paragraphs[0].add_run('중립').bold = True
        hdr_cells3[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        hdr_cells3[2].paragraphs[0].add_run('부정').bold = True
        hdr_cells3[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i in range(3):
            hdr_cells3 = table3.rows[i+1].cells
            for g in range(3):
                hdr_cells3[g].paragraphs[0].add_run(str(voice[i][g])).bold = True
                hdr_cells3[g].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER 

        p4 = document.add_paragraph()
        pp4= p4.add_run('\n녹취록\n')
        pp4.bold = True
        pp4.italic = True
        font4 = pp4.font
        font4.size = Pt(17)
        pp5= p4.add_run(self.Answer1+'\n')
        pp5= p4.add_run(self.Answer2+'\n')
        pp5= p4.add_run(self.Answer3+'\n')
        font5 = pp5.font
        font5.size = Pt(12)

        document.save('report.docx')

if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_recordWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())